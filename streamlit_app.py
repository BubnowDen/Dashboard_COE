import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime
from docx import Document
from docx.shared import Inches
import io

# Настройка страницы
st.set_page_config(page_title="Система Управления Проектами | SP 2.5", layout="wide")


# --- 1. ФУНКЦИЯ ЗАГРУЗКИ ---
@st.cache_data
def load_data(file):
    df = pd.read_excel(
        file,
        sheet_name="Svod",
        na_values=["#ЗНАЧ!", "#REF!", "#DIV/0!", "#Н/Д", " ", ""],
    )
    df.columns = [str(col).strip().replace("\n", " ") for col in df.columns]

    stages_map = {
        "1. Начало": ("Дата начала план", "Дата начала факт"),
        "2. Карточка": ("План_карточки", "Факт_карточки"),
        "3. Статистика 1": ("План_Статистика_1", "Факт_Статистика_1"),
        "4. Опрос 1": ("План опрос 1", "Факт опрос 1"),
        "5 Мероприятия": ("План_мероприятий", "Факт_мероприятий"),
        "6. Документ": ("План_документа", "Факт_документ"),
        "7. Анализ 1": ("Анализ 1 план", "Анализ 1 факт"),
        "8. Клиент 2": ("План_клиент_2", "Факт_клиент_2"),
        "9. Статистика 2": ("Статистика план 2", "Статистика факт 2"),
        "10. Анализ 2": ("Анализ 2 план", "Анализ 2 факт"),
        "11. Кейс": ("Кейс_план", "Кейс_факт"),
        "12. Защита": ("Защита план", "Защита факт"),
    }
    for stage in stages_map.values():
        p, f = stage
        if p in df.columns:
            df[p] = pd.to_datetime(df[p], errors="coerce")
        if f in df.columns:
            df[f] = pd.to_datetime(df[f], errors="coerce")

    return df, stages_map


# --- 2. ИНТЕРФЕЙС И ФИЛЬТРЫ ---
st.sidebar.header("🕹️ Навигация 2.0")
page = st.sidebar.radio(
    "Перейти к разделу:", ["📊 Аналитика (Главная)", "📑 Массовый отчет Milestone"]
)

st.sidebar.divider()
uploaded_file = st.sidebar.file_uploader("Загрузите реестр ", type=["xlsx", "xlsm"])

if uploaded_file:
    df, stages_map = load_data(uploaded_file)

    # Период анализа
    latest_plan_in_file = (
        df["Дата начала план"].max()
        if "Дата начала план" in df.columns
        else datetime.now()
    )
    default_date = (
        latest_plan_in_file
        if latest_plan_in_file.year >= datetime.now().year
        else datetime.now()
    )
    st.sidebar.subheader("📅 Период анализа")
    today_now = pd.Timestamp(st.sidebar.date_input("Дата «Сегодня»:", datetime.today()))

    # --- СЛОВАРЬ HR BP ---
    HR_BP_MAP = {
        "Динара": ["Свиноводство", "Молочное животноводство"],
        "Ольга": ["Мясопереработка"],
        "Екатерина": [
            "Производство комбикормов",
            "Плодородие",
            "Растениеводство",
            "Повышение эффективности",
        ],
        "Ирина": [
            "Закупки",
            "Логистика",
            "Качество",
            "Маркетнг",
            "Планирование",
            "ИТ",
            "Коммерческая деятельность",
        ],
        "Елена": [
            "ОТиПБ",
            "Внутренний аудит",
            "Капитальное строительство",
            "Благотворительный фонд",
            "Бухгалтерия и налоги",
            "Юридическое обеспечение",
            "Экономика и финансы",
            "PR",
            "Офисная инфраструктура",
        ],
    }
    selected_hr_bps = st.sidebar.multiselect(
        "👤 Куратор HR BP", options=list(HR_BP_MAP.keys())
    )
    bp_bl = []
    for bp in selected_hr_bps:
        bp_bl.extend(HR_BP_MAP.get(bp, []))

    # --- СОЗДАНИЕ ВСЕХ ФИЛЬТРОВ ---
    all_members = sorted(df["Член правления"].dropna().unique())
    members = st.sidebar.multiselect(
        "🏛️ Член правления", all_members, default=all_members
    )

    ba_col_name = df.columns[35] if len(df.columns) > 35 else None
    if ba_col_name:
        all_ba = sorted(df[ba_col_name].dropna().unique())
        selected_ba = st.sidebar.multiselect("🔍 Курирующий БА", all_ba, default=all_ba)

    all_bl = sorted(df["Блок"].dropna().unique())
    default_bl = [b for b in bp_bl if b in all_bl] if selected_hr_bps else all_bl
    blocks = st.sidebar.multiselect("🏢 Блок", all_bl, default=default_bl)

    all_stats = sorted(df["Статус"].dropna().unique())
    excluded_status = ["Уволен", "Отозван"]
    selected_statuses = st.sidebar.multiselect(
        "📌 Статус проекта",
        all_stats,
        default=[s for s in all_stats if s not in excluded_status],
    )

    all_fio = sorted(df["ФИО"].dropna().unique())
    selected_fio = st.sidebar.multiselect("👤 Сотрудники", all_fio, default=all_fio)

    # --- ПРИМЕНЕНИЕ МАСКИ ФИЛЬТРАЦИИ (ПОЛНЫЙ СПИСОК УСЛОВИЙ) ---
    mask = (
        df["Блок"].isin(blocks)
        & df["Член правления"].isin(members)
        & df["Статус"].isin(selected_statuses)
        & df["ФИО"].isin(selected_fio)
    )
    if ba_col_name:
        mask = mask & df[ba_col_name].isin(selected_ba)

    f_df = df[mask].copy()

    # Выбор этапов
    all_stages_list = list(stages_map.keys())
    select_all_stages = st.sidebar.checkbox("Анализировать все этапы", value=True)
    selected_stages = (
        all_stages_list
        if select_all_stages
        else st.sidebar.multiselect(
            "Этапы", all_stages_list, default=[all_stages_list[1]]
        )
    )
    # =========================================================================
    # ПРЕДВАРИТЕЛЬНЫЕ РАСЧЕТЫ (ДЛЯ ОБЕИХ СТРАНИЦ И WORD)
    # =========================================================================

    # Расчет честного KPI
    tp_s, td_s = 0, 0
    for _, r in f_df.iterrows():
        for s in selected_stages:
            pc, fc = stages_map[s]
            if pc in r and pd.notnull(r[pc]) and r[pc] <= today_now:
                tp_s += 1
                if fc in r and pd.notnull(r[fc]):
                    td_s += 1
    hkpi = round((td_s / tp_s * 100), 1) if tp_s > 0 else 100.0

    # Создание графиков (чтобы они были доступны для экспорта в Word)
    fig_pie = px.pie(f_df, names="Статус", hole=0.4, title="1. Состояние портфеля")
    fig_bar_block = px.bar(
        f_df, x="Блок", color="Статус", barmode="group", title="2. Статусы по Блокам"
    )

    # Накопительный график
    fig_line = go.Figure()
    t_p_d, t_f_d = [], []
    for s in selected_stages:
        pc, fc = stages_map[s]
        t_p_d.extend(f_df[pc].dropna().tolist())
        t_f_d.extend(f_df[fc].dropna().tolist())
    if t_p_d:
        df_p_c = (
            pd.DataFrame({"d": t_p_d, "v": 1})
            .sort_values("d")
            .groupby("d")
            .sum()
            .cumsum()
            .reset_index()
        )
        fig_line.add_trace(
            go.Scatter(
                x=df_p_c["d"],
                y=df_p_c["v"],
                name="План",
                line=dict(color="#1f77b4", width=4),
            )
        )
        df_f_r = pd.DataFrame({"d": t_f_d, "v": 1}).sort_values("d")
        if not df_f_r.empty:
            df_f_c = df_f_r.groupby("d").sum().cumsum().reset_index()
            df_f_c = df_f_c[df_f_c["d"] <= today_now]
            fig_line.add_trace(
                go.Scatter(
                    x=df_f_c["d"],
                    y=df_f_c["v"],
                    name="Факт",
                    fill="tozeroy",
                    line=dict(color="#27ae60", width=2),
                )
            )

    # Спидометр
    fig_g = go.Figure(
        go.Indicator(
            mode="gauge+number",
            value=hkpi,
            title={"text": f"Сдано: {td_s} из {tp_s}"},
            gauge={
                "axis": {"range": [0, 100]},
                "bar": {"color": "white"},
                "steps": [
                    {"range": [0, 50], "color": "#e74c3c"},
                    {"range": [50, 80], "color": "#f1c40f"},
                    {"range": [80, 100], "color": "#2ecc71"},
                ],
            },
        )
    )

    # Прогресс по этапам
    sc_data = []
    for s in all_stages_list:
        pc, fc = stages_map[s]
        fv = f_df[fc].count() if fc in f_df.columns else 0
        pv = (f_df[pc] <= today_now).sum() if pc in f_df.columns else 0
        sc_data.append(
            {"Этап": s, "Факт": fv, "План": pv, "Цвет": "green" if fv >= pv else "red"}
        )
    df_cp = pd.DataFrame(sc_data)
    fig_cp = go.Figure()
    fig_cp.add_trace(
        go.Bar(
            x=df_cp["Факт"],
            y=df_cp["Этап"],
            orientation="h",
            name="Факт (сдано)",
            marker_color=df_cp["Цвет"],
            text=df_cp["Факт"],
            textposition="inside",
        )
    )
    fig_cp.add_trace(
        go.Scatter(
            x=df_cp["План"],
            y=df_cp["Этап"],
            mode="markers+text",
            name="План",
            text=df_cp["План"],
            textposition="top right",
            marker=dict(color="#1f77b4", size=12, symbol="line-ns-open", line_width=3),
        )
    )
    fig_cp.update_layout(
        barmode="overlay", margin=dict(l=150), yaxis=dict(autorange="reversed")
    )
    # Аргумент reversed отоброжает график в обратном порядке, от Защиты до Начала а не наоборот 150 поставил что бы график не уплывал
    # =========================================================================
    # ЛОГИКА СТРАНИЦ
    # =========================================================================
    if page == "📊 Аналитика (Главная)":
        st.title(f"📊 Аналитический дашборд")
        st.info(f"Дата анализа: {today_now.strftime('%d.%m.%Y')}")

        # Ряд 1: Круги и Блоки
        c1, c2 = st.columns(2)
        with c1:
            st.plotly_chart(fig_pie, width="stretch")
        with c2:
            st.plotly_chart(fig_bar_block, width="stretch")

        # Ряд 2: Линия План/Факт
        st.divider()
        st.subheader("3. Накопительный темп сдачи")
        st.plotly_chart(fig_line, width="stretch")

        # Ряд 3: Спидометр и Точки
        st.divider()
        c3, c4 = st.columns(2)
        with c3:
            st.subheader("4. Честный индекс выполнения")
            st.plotly_chart(fig_g, width="stretch")
        with c4:
            st.subheader("5. Прогресс по этапам (Точка - План)")
            st.plotly_chart(fig_cp, width="stretch")

        # Ряд 4: Личный рейтинг
        st.divider()
        st.subheader("6. Личный рейтинг эффективности")
        fio_stat = []
        for fn in f_df["ФИО"].unique():
            sub_f = f_df[f_df["ФИО"] == fn]
            sp_f, sd_f = 0, 0
            for _, r in sub_f.iterrows():
                for s in all_stages_list:
                    pc, fc = stages_map[s]
                    if pc in r and pd.notnull(r[pc]) and r[pc] <= today_now:
                        sp_f += 1
                        if fc in r and pd.notnull(r[fc]):
                            sd_f += 1
            fio_stat.append(
                {"ФИО": fn, "КПД": round((sd_f / sp_f * 100), 1) if sp_f > 0 else 100.0}
            )
        df_fio = pd.DataFrame(fio_stat).sort_values("КПД", ascending=True)
        fig_fio = px.bar(
            df_fio,
            x="КПД",
            y="ФИО",
            orientation="h",
            text="КПД",
            color="КПД",
            color_continuous_scale="RdYlGn",
            range_color=[0, 100],
        )
        fig_fio.update_traces(textangle=0, textposition="outside", cliponaxis=False)
        st.plotly_chart(fig_fio, width="stretch")

        # Детализация Milestone (expander)
        with st.expander("🔍 Детальный аудит: Матрица Milestone"):
            sel_p = st.selectbox("Сотрудник:", options=df_fio["ФИО"].unique())
            p_sub = f_df[f_df["ФИО"] == sel_p]
            g_labs, m_vals, f_x, f_y = [], [], [], []
            for _, r in p_sub.iterrows():
                g_s = str(r["Цели по процессам"])[:45] + "..."
                g_labs.append(g_s)
                row_st = []
                for i, sn in enumerate(all_stages_list):
                    pc, fc = stages_map[sn]
                    is_p = pd.notnull(r[pc]) and r[pc] <= today_now
                    is_d = pd.notnull(r[fc])
                    if is_d:
                        s_c = 2
                    elif is_p:
                        s_c = 1
                    else:
                        s_c = 0
                    row_st.append(s_c)
                    if is_p:
                        f_x.append(i + 1)
                        f_y.append(g_s)
                m_vals.append(row_st)
            fig_h = go.Figure(
                data=go.Heatmap(
                    z=m_vals,
                    x=[f"{i + 1}" for i in range(len(all_stages_list))],
                    y=g_labs,
                    colorscale=[[0, "#34495e"], [0.5, "#e74c3c"], [1, "#27ae60"]],
                    showscale=False,
                    xgap=2,
                    ygap=2,
                    zmin=0,
                    zmax=2,
                )
            )
            fig_h.add_trace(
                go.Scatter(
                    x=f_x,
                    y=f_y,
                    mode="markers",
                    marker=dict(
                        symbol="circle", size=8, line=dict(color="#3498db", width=4)
                    ),
                    hoverinfo="skip",
                )
            )
            fig_h.update_layout(
                xaxis=dict(
                    side="top",
                    ticktext=all_stages_list,
                    tickvals=[f"{i + 1}" for i in range(len(all_stages_list))],
                ),
                yaxis=dict(autorange="reversed"),
                margin=dict(l=350),
            )
            st.plotly_chart(fig_h, width="stretch")

        # Рейтинги Блоков и Правления
        st.divider()
        st.subheader("🏆 Групповые рейтинги")

        def get_kpi_for(col):
            rk = []
            for v in f_df[col].unique():
                s_gr = f_df[f_df[col] == v]
                ps_g, ds_g = 0, 0
                for _, r in s_gr.iterrows():
                    for stg in all_stages_list:
                        pc, fc = stages_map[stg]
                        if (
                            pc in f_df.columns
                            and pd.notnull(r[pc])
                            and r[pc] <= today_now
                        ):
                            ps_g += 1
                            if fc in f_df.columns and pd.notnull(r[fc]):
                                ds_g += 1
                rk.append(
                    {
                        col: v,
                        "КПД": round((ds_g / ps_g * 100), 1) if ps_g > 0 else 100.0,
                    }
                )
            return pd.DataFrame(rk).sort_values("КПД", ascending=True)

        st.plotly_chart(
            px.bar(
                get_kpi_for("Блок"),
                x="КПД",
                y="Блок",
                orientation="h",
                text="КПД",
                color="КПД",
                color_continuous_scale="RdYlGn",
                range_color=[0, 100],
                title="Рейтинг Блоков",
            ),
            width="stretch",
        )
        st.plotly_chart(
            px.bar(
                get_kpi_for("Член правления"),
                x="КПД",
                y="Член правления",
                orientation="h",
                text="КПД",
                color="КПД",
                color_continuous_scale="RdYlGn",
                range_color=[0, 100],
                title="Рейтинг по направлениям",
            ),
            width="stretch",
        )

        # Готовность к защите
        # --- ГОТОВНОСТЬ К ЗАЩИТЕ (ВЕРСИЯ 2.1 - ДИФФЕРЕНЦИРОВАННАЯ) ---
        st.divider()
        st.subheader("🛡️ Готовность к полугодовой защите")
        ready_list = []

        # Ключи колонок из stages_map
        col_c_f = stages_map["2. Карточка"][1]
        col_d_f = stages_map["6. Документ"][1]
        col_stat_f = stages_map["3. Статистика 1"][1]
        col_surv_f = stages_map["4. Опрос 1"][1]

        # Колонки для проверки очереди документа (1-5)
        pre_cols = [
            stages_map[s][1]
            for s in [
                "1. Начало",
                "2. Карточка",
                "3. Статистика 1",
                "4. Опрос 1",
                "5 Мероприятия",
            ]
        ]

        for p in sorted(f_df["ФИО"].unique()):
            p_df = f_df[f_df["ФИО"] == p]
            pu_val = str(p_df["ПУ"].iloc[0]) if "ПУ" in p_df.columns else ""
            is_high_pu = any(x in pu_val for x in ["1", "2"])

            all_c = p_df[col_c_f].notnull().all()
            val_d = False
            for _, row in p_df.iterrows():
                if pd.notnull(row[col_d_f]) and row[pre_cols].notnull().all():
                    val_d = True
                    break

            stat_count = p_df[col_stat_f].notnull().sum()
            survey_count = p_df[col_surv_f].notnull().sum()

            if is_high_pu:
                is_r = all_c and val_d and stat_count >= 2 and survey_count >= 2
            else:
                is_r = all_c and val_d

            reason = []
            if not all_c:
                reason.append("нет всех карточек")
            if not val_d:
                reason.append("нет документа / нарушена очередь")
            if is_high_pu:
                if stat_count < 2:
                    reason.append(f"нужно 2 статистики (сейчас {stat_count})")
                if survey_count < 2:
                    reason.append(f"нужно 2 опроса (сейчас {survey_count})")

            ready_list.append(
                {
                    "Сотрудник": p,
                    "ПУ": pu_val,
                    "Статус": "✅ Готов" if is_r else "❌ Не готов",
                    "Причина": ", ".join(reason) if not is_r else "—",
                    "is_ready": is_r,
                }
            )

        df_ready = pd.DataFrame(ready_list)

        if not df_ready.empty:
            r_cnt = df_ready["is_ready"].sum()
            total_p = len(df_ready)
            st.write(
                f"📊 Всего сотрудников: **{total_p}** | Готовы к защите: **{r_cnt}**"
            )
            st.progress(r_cnt / total_p if total_p > 0 else 0)

            # Отображаем таблицу (добавили колонку ПУ для наглядности)
            st.dataframe(
                df_ready.drop(columns=["is_ready"]).style.apply(
                    lambda x: [
                        "background-color: #1b5e20; color: white"
                        if x.Статус == "✅ Готов"
                        else ""
                        for _ in x
                    ],
                    axis=1,
                ),
                width="stretch",
                hide_index=True,
            )

            # Мини-раздел: План по владельцу
            # =========================================================================
        # РАЗДЕЛ: КАЛЕНДАРЬ ДЕДЛАЙНОВ (OWNER & BA PLANNER) - SP 2.3
        # =========================================================================
        st.divider()
        with st.expander(
            "📅 Календарь ближайших дедлайнов (План на 2 месяца)", expanded=True
        ):
            ba_col_name = df.columns[35] if len(df.columns) > 35 else "Бизнес-аналитик"

            st.write(
                "Выберите аналитика, чтобы увидеть график его нагрузки на текущий и следующий месяц."
            )

            # 1. Выбор БА
            available_bas = sorted(f_df[ba_col_name].dropna().unique().tolist())
            if available_bas:
                selected_ba_planner = st.selectbox(
                    "🎯 Выберите Бизнес-аналитика для просмотра календаря:",
                    options=available_bas,
                )

                # Фильтруем данные по выбранному БА
                planner_df = f_df[f_df[ba_col_name] == selected_ba_planner].copy()

                # 2. Определяем окно планирования (Текущий месяц + следующий)
                start_window = today_now.replace(day=1)
                # Конец следующего месяца
                if today_now.month == 12:
                    end_window = today_now.replace(
                        year=today_now.year + 1, month=2, day=1
                    ) - pd.Timedelta(days=1)
                else:
                    # Приблизительно +60 дней для захвата следующего месяца целиком
                    end_window = (
                        start_window + pd.DateOffset(months=2)
                    ) - pd.Timedelta(days=1)

                # 3. Собираем все планы в одну таблицу для графика
                calendar_events = []
                plan_columns = [
                    v[0] for v in stages_map.values() if v[0] in planner_df.columns
                ]

                for _, row in planner_df.iterrows():
                    for stage_name, (p_col, f_col) in stages_map.items():
                        plan_date = row[p_col]
                        fact_date = row[f_col]

                        # Нам нужны только невыполненные планы, попадающие в наше окно
                        if pd.notnull(plan_date) and pd.isnull(fact_date):
                            if start_window <= plan_date <= end_window:
                                calendar_events.append(
                                    {
                                        "Дата": plan_date,
                                        "Цель": row["Цели по процессам"],
                                        "Владелец (бизнес)": row["ФИО"],
                                        "Этап": stage_name,
                                        "Краткая цель": (
                                            str(row["Цели по процессам"])[:40] + "..."
                                        )
                                        if len(str(row["Цели по процессам"])) > 40
                                        else row["Цели по процессам"],
                                    }
                                )

                if calendar_events:
                    df_cal = pd.DataFrame(calendar_events)

                    # 4. Строим визуальный календарь-таймлайн
                    fig_cal = px.scatter(
                        df_cal,
                        x="Дата",
                        y="Краткая цель",
                        color="Этап",
                        hover_data=["Владелец (бизнес)", "Этап", "Цель"],
                        title=f"График дедлайнов БА: {selected_ba_planner}",
                        labels={"Краткая цель": "Проект (Цель)"},
                        height=200 + (len(df_cal["Краткая цель"].unique()) * 40),
                    )

                    # Настройка маркеров и линий сетки
                    fig_cal.update_traces(marker=dict(size=12, symbol="diamond"))

                    # Выделяем "Сегодня" на календаре
                    fig_cal.add_vline(
                        x=today_now.timestamp() * 1000,
                        line_dash="dash",
                        line_color="red",
                        annotation_text="Сегодня",
                    )

                    fig_cal.update_layout(
                        xaxis=dict(
                            dtick="D1",  # Сетка по дням
                            tickformat="%d\n%b",
                            range=[start_window, end_window],
                            gridcolor="rgba(255,255,255,0.1)",
                        ),
                        yaxis=dict(autorange="reversed"),
                        legend=dict(
                            orientation="h",
                            yanchor="bottom",
                            y=1.02,
                            xanchor="right",
                            x=1,
                        ),
                    )

                    st.plotly_chart(fig_cal, width="stretch")

                    # Добавим текстовую таблицу под графиком для печати
                    st.write("📋 **Детализированный список задач:**")
                    st.dataframe(
                        df_cal[
                            ["Дата", "Этап", "Владелец (бизнес)", "Цель"]
                        ].sort_values("Дата"),
                        width="stretch",
                        hide_index=True,
                    )
                else:
                    st.info(
                        f"У аналитика {selected_ba_planner} нет запланированных дедлайнов на ближайшие 2 месяца."
                    )
            else:
                st.warning("В текущей выборке не найдено Бизнес-аналитиков.")
        # =========================================================================
        # РАЗДЕЛ: АНАЛИЗ ЗАВИСШИХ ПРОЕКТОВ (STAGNATION ALERT) - SP 2.3
        # =========================================================================
        st.divider()
        st.subheader("⚠️ Детектор простоя проектов")

        # Названия ключевых колонок
        ba_col_name = df.columns[35] if len(df.columns) > 35 else "Бизнес-аналитик"
        last_stage_fact_col = stages_map["12. Защита"][1]
        fact_cols = [v[1] for v in stages_map.values() if v[1] in f_df.columns]

        # Базовая фильтрация для анализа (убираем финалистов и мусорные статусы)
        # Статусы "Уволен" и "Отозван" уже убраны глобальным фильтром f_df,
        # но на всякий случай исключаем их и здесь + убираем завершенные.
        analysis_df = f_df[
            (f_df[last_stage_fact_col].isnull())
            & (~f_df["Статус"].isin(["Уволен", "Отозван"]))
        ].copy()

        if not analysis_df.empty:
            # Считаем дату последнего факта для каждой строки
            analysis_df["Последний_движ"] = analysis_df[fact_cols].max(axis=1)
            analysis_df["Дней_простоя"] = (
                today_now - analysis_df["Последний_движ"]
            ).dt.days

            # Разделяем на "Зависших" и "Не начатых"
            stagnant = analysis_df[analysis_df["Дней_простоя"] >= 21].copy()
            not_started = analysis_df[analysis_df["Последний_движ"].isnull()].copy()

            # Метрики сверху
            m1, m2 = st.columns(2)
            m1.metric("🔥 Зависло (21+ дней)", len(stagnant), delta_color="inverse")
            m2.metric("💤 Не начато вообще", len(not_started))

            # Вкладки для чистоты интерфейса
            tab_stagnant, tab_not_started = st.tabs(
                ["🔥 Зависшие (встали)", "💤 Не начатые (пусто в фактах)"]
            )

            with tab_stagnant:
                if not stagnant.empty:
                    st.write(
                        "Список проектов, по которым движение было, но прекратилось:"
                    )
                    st.dataframe(
                        stagnant[
                            [
                                "Блок",
                                "Цели по процессам",
                                ba_col_name,
                                "Последний_движ",
                                "Дней_простоя",
                            ]
                        ].sort_values("Дней_простоя", ascending=False),
                        column_config={
                            "Последний_движ": st.column_config.DateColumn(
                                "Дата последнего факта"
                            ),
                            "Дней_простоя": st.column_config.NumberColumn(
                                "Дней без движения", format="%d дн. ⏳"
                            ),
                            "Цели по процессам": st.column_config.TextColumn(
                                "Цель", width="large"
                            ),
                        },
                        width="stretch",
                        hide_index=True,
                    )

                    # Маленький график - кто из БА "накопил" больше всего простоев
                    fig_stag_ba = px.bar(
                        stagnant.groupby(ba_col_name).size().reset_index(name="Кол-во"),
                        x=ba_col_name,
                        y="Кол-во",
                        title="Застои в разрезе БА",
                        color_discrete_sequence=["#e74c3c"],
                    )
                    st.plotly_chart(fig_stag_ba, width="stretch")
                else:
                    st.success("Активных зависших проектов не обнаружено.")

            with tab_not_started:
                if not not_started.empty:
                    st.write(
                        "Проекты, по которым не проставлено ни одного факта (даже старт):"
                    )
                    st.dataframe(
                        not_started[
                            ["Блок", "Цели по процессам", ba_col_name, "Статус"]
                        ],
                        width="stretch",
                        hide_index=True,
                    )
                else:
                    st.success("Все проекты имеют хотя бы один зафиксированный факт!")
        else:
            st.info("Нет активных проектов для анализа застоя.")

        # =========================================================================
        # РАЗДЕЛ: АНАЛИТИКА ЗАГРУЗКИ БИЗНЕС-АНАЛИТИКОВ (SP 2.2 - ИСПРАВЛЕННЫЙ)
        # =========================================================================
        st.divider()
        st.subheader("⚖️ Анализ нагрузки Бизнес-аналитиков")

        # Явно определяем названия колонок
        # AJ - это 36-я колонка (индекс 35). Проверим, как она называется в загруженном файле
        ba_col_name = df.columns[35] if len(df.columns) > 35 else "Бизнес-аналитик"
        fio_col_name = "ФИО"  # Владелец процесса

        # 1. ТАБЛИЦА ТЕКУЩЕЙ ЗАГРУЗКИ (ГРУППИРОВКА ПО БА)
        with st.expander("📊 Реальная загрузка БА (на основе Excel)", expanded=False):
            if not f_df.empty:
                # Группируем СТРОГО по Бизнес-аналитику
                ba_summary = (
                    f_df.groupby(ba_col_name)
                    .agg(
                        {
                            "Цели по процессам": "count",  # Кол-во строк (целей)
                            fio_col_name: "nunique",  # Кол-во уникальных ФИО владельцев
                        }
                    )
                    .reset_index()
                )

                ba_summary.columns = [
                    "Бизнес-аналитик",
                    "Кол-во целей",
                    "Кол-во владельцев процессов (ФИО)",
                ]

                # Сортируем по нагрузке (у кого больше целей)
                st.dataframe(
                    ba_summary.sort_values("Кол-во целей", ascending=False),
                    width="stretch",
                    hide_index=True,
                )
            else:
                st.info("Данные не найдены. Проверьте фильтры.")

                # 2. ПЕСОЧНИЦА МОДЕЛИРОВАНИЯ (ГРУППИРОВКА ПО БА)
        with st.expander(
            "🏗️ Песочница: Перераспределение нагрузки между БА", expanded=False
        ):
            st.info(
                "Нажмите на заголовок столбца 'Блок' для быстрой сортировки. Измените имя БА, чтобы увидеть прогноз."
            )

            if not f_df.empty:
                # Берем список всех БА из ИСХОДНОГО файла
                list_of_all_ba = sorted(df[ba_col_name].dropna().unique().tolist())

                # ДОБАВИЛИ 'Блок' в таблицу Песочницы
                sandbox_df = f_df[
                    ["Блок", "Цели по процессам", fio_col_name, ba_col_name]
                ].copy()

                edited_df = st.data_editor(
                    sandbox_df,
                    column_config={
                        "Блок": st.column_config.TextColumn(
                            "Блок", disabled=True, width="small"
                        ),
                        ba_col_name: st.column_config.SelectboxColumn(
                            "Курирующий БА (изменить)",
                            options=list_of_all_ba,
                            width="large",
                            required=True,
                        ),
                        "Цели по процессам": st.column_config.TextColumn(
                            "Цель", disabled=True
                        ),
                        fio_col_name: st.column_config.TextColumn(
                            "Владелец (бизнес)", disabled=True
                        ),
                    },
                    hide_index=True,
                    width="stretch",
                    key="sandbox_ba_editor_v2",  # Обновили ключ для сброса состояния
                )

                # --- РАСЧЕТ ПРОГНОЗА (логика остается прежней) ---
                st.write("---")
                st.write("### 📈 Прогноз нагрузки (после ваших правок)")

                forecast_ba = (
                    edited_df.groupby(ba_col_name)
                    .agg({"Цели по процессам": "count", fio_col_name: "nunique"})
                    .reset_index()
                )
                forecast_ba.columns = ["БА", "Прогноз целей", "Прогноз владельцев"]

                fig_forecast = go.Figure()
                fig_forecast.add_trace(
                    go.Bar(
                        x=forecast_ba["БА"],
                        y=forecast_ba["Прогноз целей"],
                        name="Кол-во целей",
                        marker_color="#1f77b4",
                        text=forecast_ba["Прогноз целей"],
                        textposition="auto",
                    )
                )
                fig_forecast.add_trace(
                    go.Scatter(
                        x=forecast_ba["БА"],
                        y=forecast_ba["Прогноз владельцев"],
                        name="Уникальных владельцев (ФИО)",
                        mode="lines+markers+text",
                        text=forecast_ba["Прогноз владельцев"],
                        textposition="top center",
                        line=dict(color="#e74c3c", width=3),
                        yaxis="y2",
                    )
                )
                fig_forecast.update_layout(
                    title="Сравнение нагрузки БА в модели",
                    yaxis=dict(title="Кол-во целей"),
                    yaxis2=dict(
                        title="Кол-во владельцев (ФИО)", overlaying="y", side="right"
                    ),
                    legend=dict(
                        orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1
                    ),
                )
                st.plotly_chart(fig_forecast, width="stretch")

                if st.button("Сбросить модель"):
                    st.rerun()
            else:
                st.warning("Выберите данные в фильтрах для начала моделирования.")
    # =========================================================================
    # СТРАНИЦА 2: МАССОВЫЙ ОТЧЕТ
    # =========================================================================
    elif page == "📑 Массовый отчет Milestone":
        st.title("📑 Массовая выгрузка карт Milestone")
        if not f_df.empty:
            for person in sorted(f_df["ФИО"].unique()):
                st.markdown(f"---")
                p_sub = f_df[f_df["ФИО"] == person]
                st.subheader(f"👤 {person} ({p_sub['Блок'].iloc[0]})")
                g_labs, m_vals, t_vals, f_x, f_y = [], [], [], [], []
                for _, row in p_sub.iterrows():
                    g_full = str(row["Цели по процессам"])
                    g_short = (g_full[:65] + "..") if len(g_full) > 65 else g_full
                    g_labs.append(g_short)
                    row_st, row_tx = [], []
                    for i, sn in enumerate(all_stages_list):
                        pc, fc = stages_map[sn]
                        val_p, val_f = row[pc], row[fc]
                        is_p_exists = pd.notnull(val_p)
                        is_p_pass = is_p_exists and val_p <= today_now
                        is_d = pd.notnull(val_f)
                        if is_d:
                            state, txt = 2, ""
                        elif is_p_pass:
                            state, txt = 1, ""
                        else:
                            state, txt = (
                                0,
                                val_p.strftime("%d.%m") if is_p_exists else "",
                            )
                        row_st.append(state)
                        row_tx.append(txt)
                        if is_p_pass:
                            f_x.append(i + 1)
                            f_y.append(g_short)
                    m_vals.append(row_st)
                    t_vals.append(row_tx)
                fig_h = go.Figure(
                    data=go.Heatmap(
                        z=m_vals,
                        x=[f"{i + 1}" for i in range(len(all_stages_list))],
                        y=g_labs,
                        text=t_vals,
                        texttemplate="%{text}",
                        textfont={"size": 10},
                        colorscale=[[0, "#34495e"], [0.5, "#e74c3c"], [1, "#27ae60"]],
                        showscale=False,
                        xgap=2,
                        ygap=2,
                        zmin=0,
                        zmax=2,
                    )
                )
                fig_h.add_trace(
                    go.Scatter(
                        x=f_x,
                        y=f_y,
                        mode="markers",
                        marker=dict(
                            symbol="circle",
                            size=8,
                            line=dict(color="#3498db", width=2.5),
                        ),
                        hoverinfo="skip",
                    )
                )
                fig_h.update_layout(
                    xaxis=dict(
                        side="top",
                        ticktext=all_stages_list,
                        tickvals=[f"{i + 1}" for i in range(len(all_stages_list))],
                        tickangle=-45,
                    ),
                    yaxis=dict(autorange="reversed"),
                    height=180 + (len(p_sub) * 45),
                    margin=dict(l=450, t=80),
                )
                st.plotly_chart(fig_h, width="stretch", key=f"bulk_{person}")

    # =========================================================================
    # ЭКСПОРТ В WORD (Доступен всегда)
    # =========================================================================
    st.sidebar.divider()
    if st.sidebar.button("Сгенерировать отчет в Word"):
        try:
            with st.spinner("Создаю документ..."):
                doc = Document()
                doc.add_heading(f"Отчет на {today_now.strftime('%d.%m.%Y')}", 0)
                img_bytes = fig_g.to_image(format="png")
                doc.add_picture(io.BytesIO(img_bytes), width=Inches(6))
                doc_io = io.BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)
                st.sidebar.download_button(
                    label="⬇️ Скачать Word", data=doc_io, file_name="Report.docx"
                )
        except Exception as e:
            st.sidebar.error(f"Ошибка экспорта: {e}")

    st.divider()
    st.markdown(
        "<div style='text-align: center; color: gray; font-style: italic;'>Система SP 2.1 | Бубнов Д.А | СОЭ | 2026.</div>",
        unsafe_allow_html=True,
    )
else:
    st.info("Загрузите реестр (Dashboard.xlsm) для начала работы.")
